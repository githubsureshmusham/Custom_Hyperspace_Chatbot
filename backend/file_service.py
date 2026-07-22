"""Extract readable text (or image data) from uploaded files so the LLM can use it.

Supported:
- Plain text / source code / markdown / json / csv / logs, etc.
- PDF  (pypdf)
- DOCX (python-docx)
- XLSX (openpyxl)
- Images (returned as base64 data URLs for vision-capable models)

Anything else is attempted as UTF-8 text; if that fails we report the file as
binary/unsupported rather than crashing.
"""
from __future__ import annotations

import base64
import csv
import io
import os

# Max characters of extracted text we keep per file, to avoid blowing the
# model's context window. ~48k chars ≈ 12k tokens.
MAX_TEXT_CHARS = 48_000

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".yaml", ".yml",
    ".xml", ".html", ".htm", ".log", ".ini", ".cfg", ".conf", ".env",
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".c", ".cpp", ".h", ".hpp",
    ".cs", ".go", ".rs", ".rb", ".php", ".sh", ".bat", ".ps1", ".sql",
    ".css", ".scss", ".less", ".r", ".kt", ".swift", ".dart", ".vue",
    ".toml", ".gradle", ".properties",
}

IMAGE_MIME = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
}


def _truncate(text: str) -> str:
    if len(text) > MAX_TEXT_CHARS:
        return text[:MAX_TEXT_CHARS] + "\n\n…[content truncated]…"
    return text


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[Cannot read PDF: 'pypdf' is not installed. Run: pip install pypdf]"

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            txt = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            txt = ""
        if txt.strip():
            parts.append(f"--- Page {i} ---\n{txt}")
    return "\n\n".join(parts) if parts else "[No extractable text in PDF]"


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        return "[Cannot read DOCX: 'python-docx' is not installed. Run: pip install python-docx]"

    document = docx.Document(io.BytesIO(data))
    lines = [p.text for p in document.paragraphs if p.text.strip()]
    # Include tables too
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines) if lines else "[No text in DOCX]"


def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "[Cannot read XLSX: 'openpyxl' is not installed. Run: pip install openpyxl]"

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"--- Sheet: {ws.title} ---")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                parts.append(", ".join(cells))
    wb.close()
    return "\n".join(parts) if parts else "[No data in spreadsheet]"


def _extract_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    try:
        reader = csv.reader(io.StringIO(text))
        rows = [", ".join(r) for r in reader]
        return "\n".join(rows)
    except Exception:  # noqa: BLE001
        return text


def extract_file(filename: str, data: bytes) -> dict:
    """Return a dict describing the extracted content.

    Shapes:
      {"type": "text",  "filename": ..., "content": "<text>"}
      {"type": "image", "filename": ..., "data_url": "data:image/png;base64,..."}
      {"type": "unsupported", "filename": ..., "reason": "..."}
    """
    ext = os.path.splitext(filename)[1].lower()

    # Images -> base64 data URL for vision models
    if ext in IMAGE_EXTS:
        mime = IMAGE_MIME.get(ext, "image/png")
        b64 = base64.b64encode(data).decode("ascii")
        return {
            "type": "image",
            "filename": filename,
            "data_url": f"data:{mime};base64,{b64}",
        }

    try:
        if ext == ".pdf":
            content = _extract_pdf(data)
        elif ext == ".docx":
            content = _extract_docx(data)
        elif ext == ".xlsx":
            content = _extract_xlsx(data)
        elif ext in (".csv", ".tsv"):
            content = _extract_csv(data)
        elif ext in TEXT_EXTS or ext == "":
            content = data.decode("utf-8", errors="replace")
        else:
            # Last resort: try UTF-8; if it's mostly binary, mark unsupported.
            decoded = data.decode("utf-8", errors="replace")
            # Heuristic: too many replacement chars -> binary
            if decoded.count("\ufffd") > max(20, len(decoded) * 0.05):
                return {
                    "type": "unsupported",
                    "filename": filename,
                    "reason": f"Binary or unsupported file type ({ext or 'no ext'}).",
                }
            content = decoded
    except Exception as exc:  # noqa: BLE001
        return {
            "type": "unsupported",
            "filename": filename,
            "reason": f"Could not read file: {exc}",
        }

    return {"type": "text", "filename": filename, "content": _truncate(content)}